import os
import re
import uuid
from pathlib import Path
from urllib.parse import urlparse, parse_qsl, urlencode, urlunparse

import requests
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from openai import OpenAI
from openpyxl import Workbook
from pydantic import BaseModel
from psycopg_pool import ConnectionPool
from psycopg.rows import dict_row
from pgvector.psycopg import register_vector
from docx import Document

load_dotenv()

OPENAI_API_KEY = os.environ["OPENAI_API_KEY"].strip()
DATABASE_URL = os.environ["DATABASE_URL"].strip()

EMBED_MODEL = os.getenv("EMBED_MODEL", "text-embedding-3-small")
CHAT_MODEL = os.getenv("CHAT_MODEL", "gpt-4.1-mini")
TOP_K = int(os.getenv("TOP_K", "6"))


def _configure_conn(conn):
    register_vector(conn)


pool = ConnectionPool(
    conninfo=DATABASE_URL,
    min_size=1,
    max_size=5,
    kwargs={"row_factory": dict_row},
    configure=_configure_conn,
    open=True,
)
client = OpenAI(api_key=OPENAI_API_KEY)

app = FastAPI()

# CORS (solo una vez)
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://maxconsulta.com",
        "https://www.maxconsulta.com",
        "https://app.maxconsulta.com",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

DOCS_DIR = Path("/tmp/maxconsulta_docs")
SHEETS_DIR = Path("/tmp/maxconsulta_sheets")
DOCS_DIR.mkdir(parents=True, exist_ok=True)
SHEETS_DIR.mkdir(parents=True, exist_ok=True)

ALLOWLIST = {
    "boe.es", "www.boe.es",
    "sepe.es", "www.sepe.es",
    "seg-social.es", "www.seg-social.es",
    "borm.es", "www.borm.es",
    "bocm.es", "www.bocm.es",
    "comunidad.madrid", "www.comunidad.madrid",
}

RAG_SIM_THRESHOLD = 0.72


class AskIn(BaseModel):
    question: str
    collection: str | None = None
    top_k: int | None = None


class ChatMessage(BaseModel):
    role: str
    content: str


class ChatIn(BaseModel):
    message: str | None = None
    question: str | None = None
    history: list[ChatMessage] = []
    collection: str | None = None
    area: str | None = None
    tone: str | None = "neutro"
    top_k: int | None = 6
    internet: bool = True
    url: str | None = None


class DocumentGenerateIn(BaseModel):
    template_id: str
    data: dict = {}
    draft_text: str | None = None


class SpreadsheetGenerateIn(BaseModel):
    template_id: str
    data: dict = {}
    draft_text: str | None = None


REQUIRED_BY_TEMPLATE = {
    "carta_despido_objetivo": [
        "empresa_nombre", "empresa_cif", "empresa_domicilio",
        "trabajador_nombre", "trabajador_dni", "fecha_efectos", "causa", "localidad_fecha_firma",
    ],
    "gastos_autonomo_basico": ["periodo"],
}


def _render_docx_carta_despido(data: dict, draft_text: str | None) -> Document:
    doc = Document()
    doc.add_heading("Carta de despido objetivo", level=1)
    doc.add_paragraph(f"Empresa: {data.get('empresa_nombre', '')}")
    doc.add_paragraph(f"CIF: {data.get('empresa_cif', '')}")
    doc.add_paragraph(f"Domicilio: {data.get('empresa_domicilio', '')}")
    doc.add_paragraph("")
    doc.add_paragraph(f"Trabajador/a: {data.get('trabajador_nombre', '')}")
    doc.add_paragraph(f"DNI/NIE: {data.get('trabajador_dni', '')}")
    doc.add_paragraph("")
    doc.add_paragraph(f"Fecha de efectos: {data.get('fecha_efectos', '')}")
    doc.add_paragraph("Motivo / causa:")
    doc.add_paragraph(data.get("causa", ""))
    doc.add_paragraph("")
    doc.add_paragraph(f"Localidad y fecha: {data.get('localidad_fecha_firma', '')}")
    doc.add_paragraph("Firma: ______________________________")
    if draft_text:
        doc.add_paragraph("")
        doc.add_heading("Borrador", level=2)
        doc.add_paragraph(draft_text)
    return doc


def _render_docx_reclamacion_previa_ss(data: dict, draft_text: str | None) -> Document:
    doc = Document()
    doc.add_heading("Reclamación previa Seguridad Social", level=1)
    doc.add_paragraph(f"Interesado/a: {data.get('interesado_nombre', '')}")
    doc.add_paragraph(f"DNI/NIE: {data.get('interesado_dni', '')}")
    doc.add_paragraph(f"Expediente: {data.get('expediente', '')}")
    doc.add_paragraph(f"Hechos: {data.get('hechos', '')}")
    doc.add_paragraph(f"Solicitud: {data.get('solicitud', '')}")
    if draft_text:
        doc.add_paragraph("")
        doc.add_heading("Borrador", level=2)
        doc.add_paragraph(draft_text)
    return doc


def _render_docx_solicitud_generica(data: dict, draft_text: str | None) -> Document:
    doc = Document()
    doc.add_heading("Solicitud genérica", level=1)
    doc.add_paragraph(f"Organismo destinatario: {data.get('organismo', '')}")
    doc.add_paragraph(f"Solicitante: {data.get('solicitante_nombre', '')}")
    doc.add_paragraph(f"DNI/NIE: {data.get('solicitante_dni', '')}")
    doc.add_paragraph(f"Asunto: {data.get('asunto', '')}")
    doc.add_paragraph(f"Exposición: {data.get('exposicion', '')}")
    doc.add_paragraph(f"Solicitud concreta: {data.get('peticion', '')}")
    if draft_text:
        doc.add_paragraph("")
        doc.add_heading("Borrador", level=2)
        doc.add_paragraph(draft_text)
    return doc


def _render_docx_generico(data: dict, draft_text: str | None) -> Document:
    doc = Document()
    title = (data.get("title") or "").strip()
    body = (data.get("body") or "").strip()
    footer = (data.get("footer") or "").strip()

    if title:
        doc.add_heading(title, level=1)

    chunks = [p.strip() for p in body.split("\n\n") if p.strip()]
    if not chunks and draft_text:
        chunks = [p.strip() for p in (draft_text or "").split("\n\n") if p.strip()]

    for c in chunks:
        doc.add_paragraph(c)

    if footer:
        doc.add_paragraph("")
        doc.add_paragraph(footer)

    return doc


def _render_xlsx_gastos_autonomo(data: dict, rows: list[dict] | None) -> Workbook:
    wb = Workbook()
    ws = wb.active
    ws.title = "Gastos"
    ws.append(["Periodo", data.get("periodo", "")])
    ws.append([])
    ws.append(["Fecha", "Concepto", "Importe"])

    total = 0.0
    for r in rows or []:
        importe = float(r.get("importe", 0) or 0)
        total += importe
        ws.append([r.get("fecha", ""), r.get("concepto", ""), importe])

    ws.append([])
    ws.append(["TOTAL", "", total])
    return wb


def _render_xlsx_registro_horas(data: dict, rows: list[dict] | None) -> Workbook:
    wb = Workbook()
    ws = wb.active
    ws.title = "Registro horas"
    ws.append(["Empleado", data.get("empleado", "")])
    ws.append([])
    ws.append(["Fecha", "Entrada", "Salida", "Horas"])
    for r in rows or []:
        ws.append([r.get("fecha", ""), r.get("entrada", ""), r.get("salida", ""), r.get("horas", "")])
    return wb


DOCX_TEMPLATES = {
    "carta_despido_objetivo": {
        "id": "carta_despido_objetivo",
        "title": "Carta de despido objetivo (plantilla)",
        "description": "Carta laboral básica de despido objetivo.",
        "required_fields": REQUIRED_BY_TEMPLATE["carta_despido_objetivo"],
        "optional_fields": ["referencias_legales", "cargo_firmante"],
        "render": _render_docx_carta_despido,
    },
    "reclamacion_previa_ss": {
        "id": "reclamacion_previa_ss",
        "title": "Reclamación previa SS",
        "description": "Escrito base de reclamación previa en Seguridad Social.",
        "required_fields": ["interesado_nombre", "interesado_dni", "hechos", "solicitud"],
        "optional_fields": ["expediente"],
        "render": _render_docx_reclamacion_previa_ss,
    },
    "solicitud_generica": {
        "id": "solicitud_generica",
        "title": "Solicitud genérica",
        "description": "Plantilla universal para administración.",
        "required_fields": ["organismo", "solicitante_nombre", "solicitante_dni", "peticion"],
        "optional_fields": ["asunto", "exposicion"],
        "render": _render_docx_solicitud_generica,
    },
    "doc_generico": {
        "id": "doc_generico",
        "title": "Documento Word (genérico)",
        "description": "Exporta cualquier respuesta textual a DOCX.",
        "required_fields": ["body"],
        "optional_fields": ["title", "footer"],
        "render": _render_docx_generico,
    },
}

XLSX_TEMPLATES = {
    "gastos_autonomo_basico": {
        "id": "gastos_autonomo_basico",
        "title": "Gastos Autónomo Básico (plantilla)",
        "description": "Libro simple de gastos con total.",
        "required_fields": REQUIRED_BY_TEMPLATE["gastos_autonomo_basico"],
        "optional_fields": ["rows"],
        "render": _render_xlsx_gastos_autonomo,
    },
    "registro_horas": {
        "id": "registro_horas",
        "title": "Registro de horas",
        "description": "Timesheet básico por día.",
        "required_fields": ["empleado"],
        "optional_fields": ["rows"],
        "render": _render_xlsx_registro_horas,
    },
}


@app.get("/")
def root():
    return {"ok": True, "service": "api"}


@app.get("/health")
def health():
    return {"ok": True}


def embed_query(text: str):
    # embeddings API: input puede ser string o lista; aquí lo dejamos simple
    resp = client.embeddings.create(model=EMBED_MODEL, input=text)
    return resp.data[0].embedding


def retrieve_chunks(query_embedding, match_count: int, collection_name: str | None):
    """
    Búsqueda por similitud vectorial en Postgres (pgvector).
    Llama a la función SQL match_chunks(query_embedding, match_count,
    p_collection_name, p_storage_bucket). p_storage_bucket se deja en None.
    Devuelve lista de dicts con: id, document_id, title, content,
    page_start, page_end, similarity.
    """
    import numpy as np

    vec = np.array(query_embedding, dtype=np.float32)

    with pool.connection() as conn:
        rows = conn.execute(
            "SELECT * FROM match_chunks(%s, %s, %s, %s)",
            (vec, match_count, collection_name, None),
        ).fetchall()

    return rows or []


@app.post("/ask")
def ask(inp: AskIn):
    q = (inp.question or "").strip()
    if not q:
        return {"answer": "Escribe una pregunta.", "sources": []}

    top_k = inp.top_k or TOP_K

    try:
        q_emb = embed_query(q)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generando embedding: {e}")

    try:
        rows = retrieve_chunks(q_emb, top_k, inp.collection)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    context_blocks = []
    sources = []

    for r in rows:
        title = r.get("title") or "documento"
        p1 = r.get("page_start")
        p2 = r.get("page_end")
        content = (r.get("content") or "").strip()

        sources.append({
            "title": title,
            "page_start": p1,
            "page_end": p2,
            "similarity": float(r.get("similarity", 0) or 0),
        })

        cite = f"[{title}, pág. {p1}{'' if p1 == p2 else '-' + str(p2)}]"
        context_blocks.append(f"{cite}\n{content}")

    context = "\n\n---\n\n".join(context_blocks)

    system = (
        "Eres un asistente experto en normativa y gestión de gestoría (laboral, fiscal y seguridad social). "
        "Responde en español claro. "
        "Usa SOLO la información del CONTEXTO. "
        "Si falta información, dilo. "
        "Incluye citas al final de frases relevantes usando el formato [titulo, pág. x]."
    )

    user = f"PREGUNTA:\n{q}\n\nCONTEXTO:\n{context}"

    try:
        resp = client.chat.completions.create(
            model=CHAT_MODEL,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
            temperature=0.2,
        )
        answer = resp.choices[0].message.content.strip()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error en chat completion: {e}")

    return {"answer": answer, "sources": sources}


def _best_similarity(sources: list[dict]) -> float:
    vals = []
    for s in sources or []:
        sim = s.get("similarity")
        try:
            vals.append(float(sim))
        except (TypeError, ValueError):
            continue
    return max(vals) if vals else 0.0


def _extract_url(text: str) -> str | None:
    m = re.search(r"https?://[^\s)]+", text or "")
    return m.group(0) if m else None


def _detect_official_id(text: str) -> tuple[str, str] | None:
    t = text or ""
    m_boe = re.search(r"\b(BOE-[A-Z]-\d{4}-\d{3,6})\b", t, flags=re.IGNORECASE)
    if m_boe:
        return ("boe", m_boe.group(1).upper())
    m_borme = re.search(r"\b(BORME-[A-Z]-\d{4}-\d{3,6})\b", t, flags=re.IGNORECASE)
    if m_borme:
        return ("borme", m_borme.group(1).upper())
    return None


def _search_borme_id(id_value: str) -> tuple[str | None, str]:
    api_key = os.getenv("SERPAPI_API_KEY", "").strip()
    if not api_key:
        return (None, "search_provider_missing")
    q = f"site:boe.es/borme {id_value}"
    try:
        resp = requests.get(
            "https://serpapi.com/search.json",
            params={"engine": "google", "q": q, "api_key": api_key, "num": 3},
            timeout=12,
            headers={"User-Agent": "MaxConsultaBot/1.0"},
        )
        resp.raise_for_status()
        data = resp.json()
        for item in data.get("organic_results", [])[:3]:
            link = (item.get("link") or "").strip()
            if link and _is_allowed_url(link) and "/borme" in link:
                return (link, "direct_official_id")
        return (None, "no_official_results")
    except Exception:
        return (None, "search_provider_error")


def _is_allowed_url(url: str) -> bool:
    host = (urlparse(url).netloc or "").lower()
    return host in ALLOWLIST


def _clean_text_from_html(html: str) -> str:
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript", "header", "footer", "nav"]):
        tag.decompose()
    return " ".join(soup.stripped_strings)


def _internet_fetch(url: str, question: str) -> tuple[str, list[dict]]:
    headers = {"User-Agent": "MaxConsultaBot/1.0"}
    r = requests.get(url, timeout=10, headers=headers)
    r.raise_for_status()
    soup = BeautifulSoup(r.text, "html.parser")
    title = (soup.title.string.strip() if soup.title and soup.title.string else url)
    text = _clean_text_from_html(r.text)
    snippet = text[:200] + ("..." if len(text) > 200 else "")
    answer = (
        f"He consultado la fuente oficial permitida: {title}. "
        f"Resumen rápido relacionado con tu consulta '{question}':\n\n{snippet}"
    )
    return answer, [{"title": title, "url": url, "snippet": snippet}]


def _normalize_url_for_dedup(url: str) -> str:
    try:
        u = urlparse(url)
        query = [(k, v) for k, v in parse_qsl(u.query, keep_blank_values=True) if not k.lower().startswith("utm_")]
        clean = u._replace(query=urlencode(query), fragment="")
        return urlunparse(clean)
    except Exception:
        return url


def _score_official_url(url: str) -> int:
    u = (url or "").lower()
    score = 0
    # SEPE ranking deterministic
    if "sepe.es" in u or "sede.sepe.gob.es" in u:
        if "/homesepe/personas/" in u:
            score += 6
        if "prestaciones" in u or "prestaciones-por-desempleo" in u:
            score += 4
        if "subsidios" in u:
            score += 4
        if "requisitos" in u:
            score += 3
        if "/preguntas-frecuentes/" in u or "/faqs/" in u or "detalle-pregunta" in u or "?detail=" in u:
            score -= 8
        if "cookies" in u or "privacidad" in u:
            score -= 5
    return score


def _search_official_sources(query: str, area: str | None = None) -> tuple[str, list[dict], str, list[str], list[dict]]:
    api_key = os.getenv("SERPAPI_API_KEY", "").strip()
    if not api_key:
        return (
            "Activa SERPAPI_API_KEY en el servidor para buscar en fuentes oficiales.",
            [],
            "search_provider_missing",
            [],
            [],
        )

    merged_query = query if not area else f"{query} {area}"
    official_scope = (
        "(site:boe.es OR site:www.boe.es OR site:sepe.es OR site:www.sepe.es OR site:sede.sepe.gob.es OR "
        "site:seg-social.es OR site:www.seg-social.es OR site:borm.es OR site:www.borm.es OR "
        "site:bocm.es OR site:www.bocm.es OR site:comunidad.madrid OR site:www.comunidad.madrid)"
    )

    sepe_queries = [
        "site:sepe.es subsidio por desempleo requisitos",
        "site:sepe.es HomeSepe Personas prestaciones subsidios requisitos",
        'site:sepe.es "subsidios" "requisitos"',
    ]
    queries = [f"{official_scope} {merged_query}"] + sepe_queries

    raw_urls: list[str] = []
    try:
        for q in queries:
            resp = requests.get(
                "https://serpapi.com/search.json",
                params={"engine": "google", "q": q, "api_key": api_key, "num": 5},
                timeout=12,
                headers={"User-Agent": "MaxConsultaBot/1.0"},
            )
            resp.raise_for_status()
            data = resp.json()
            for item in data.get("organic_results", [])[:5]:
                link = (item.get("link") or "").strip()
                if link and _is_allowed_url(link):
                    raw_urls.append(link)
    except Exception:
        return ("Error al consultar proveedor de búsqueda oficial.", [], "search_provider_error", [], [])

    # dedup normalized preserving first seen original
    normalized_seen = set()
    dedup_urls = []
    for u in raw_urls:
        nu = _normalize_url_for_dedup(u)
        if nu not in normalized_seen:
            normalized_seen.add(nu)
            dedup_urls.append(nu)

    scored_urls = sorted(dedup_urls, key=lambda x: _score_official_url(x), reverse=True)
    urls = scored_urls[:3]

    # fallback to original dedup top3 if scoring wiped list (defensive)
    if not urls:
        urls = dedup_urls[:3]

    if not urls:
        return (
            "No encontré resultados oficiales relevantes. Prueba afinando la consulta.",
            [],
            "no_official_results",
            [],
            [],
        )

    sources: list[dict] = []
    snippets: list[str] = []
    citations: list[dict] = []

    for i, u in enumerate(urls, start=1):
        try:
            headers = {"User-Agent": "MaxConsultaBot/1.0"}
            r = requests.get(u, timeout=10, headers=headers)
            r.raise_for_status()
            content_type = (r.headers.get("content-type") or "").lower()
            title = u
            snippet = "Fuente oficial consultada"

            if "pdf" in content_type or u.lower().endswith(".pdf"):
                snippet = "Documento PDF oficial (sin volcado de contenido)."
            else:
                soup = BeautifulSoup(r.text, "html.parser")
                title = (soup.title.string.strip() if soup.title and soup.title.string else u)
                text = _clean_text_from_html(r.text)[:12000]
                snippet = text[:200] + ("..." if len(text) > 200 else "")

            sources.append({"title": title, "url": u, "snippet": snippet})
            citations.append({"n": i, "url": u})
            snippets.append(f"[{i}] {title}: {snippet}")
        except Exception:
            continue

    if not sources:
        return (
            "No pude recuperar contenido útil de los resultados oficiales.",
            [],
            "no_official_results",
            urls,
            [],
        )

    urls_text = "\n".join([f"[{i}] {u}" for i, u in enumerate([s['url'] for s in sources], start=1)])
    answer = (
        "Resumen con fuentes oficiales consultadas:\n\n"
        + "\n".join(snippets)
        + "\n\nFuentes oficiales consultadas:\n"
        + urls_text
    )
    return (answer, sources, "searched_official_sources", [s["url"] for s in sources], citations)


def _detect_template(message: str) -> tuple[str | None, str | None]:
    t = (message or "").lower()
    if "despido" in t or "carta de despido" in t:
        return "docx", "carta_despido_objetivo"
    if "reclamación" in t or "reclamacion" in t or "seguridad social" in t or "inss" in t:
        return "docx", "reclamacion_previa_ss"
    if "solicitud" in t or "instancia" in t or "escrito" in t:
        return "docx", "solicitud_generica"
    if "excel" in t and ("gasto" in t or "autonom" in t):
        return "xlsx", "gastos_autonomo_basico"
    if "timesheet" in t or "registro" in t or "horas" in t or "jornada" in t:
        return "xlsx", "registro_horas"
    return None, None


def _extract_fields(history: list[ChatMessage], message: str) -> dict:
    text = "\n".join([m.content for m in history if m.role == "user"] + [message])
    out = {}
    patterns = {
        "empresa_nombre": r"empresa\s*[:\-]\s*([^\n,;]+)",
        "empresa_cif": r"cif\s*[:\-]\s*([A-Z0-9\-]+)",
        "empresa_domicilio": r"domicilio\s*[:\-]\s*([^\n;]+)",
        "trabajador_nombre": r"trabajador\s*[:\-]\s*([^\n,;]+)",
        "trabajador_dni": r"dni\s*[:\-]\s*([A-Z0-9\-]+)",
        "fecha_efectos": r"fecha de efectos\s*[:\-]\s*([^\n,;]+)",
        "causa": r"causa\s*[:\-]\s*([^\n]+)",
        "localidad_fecha_firma": r"localidad y fecha\s*[:\-]\s*([^\n,;]+)",
        "periodo": r"periodo\s*[:\-]\s*([^\n,;]+)",
        "empleado": r"empleado\s*[:\-]\s*([^\n,;]+)",
    }
    for key, pat in patterns.items():
        m = re.search(pat, text, flags=re.IGNORECASE)
        if m:
            out[key] = m.group(1).strip()
    return out


@app.post("/chat")
def chat(inp: ChatIn):
    question = (inp.message or inp.question or "").strip()
    if not question:
        return {
            "answer": "Escribe una pregunta.",
            "sources": [],
            "rag_sufficient": False,
            "internet_used": False,
            "internet_reason": "internet_disabled",
            "internet_sources": [],
            "needs_data": False,
            "fields_required": [],
            "can_generate_doc": False,
            "doc_type": None,
            "extracted_fields": {},
            "draft_text": None,
            "download_url": None,
            "url_used": None,
            "citations": [],
            "internet_effective": False,
            "auto_official_triggered": False,
            "rag_override_reason": None,
        }

    history_lines = []
    for h in (inp.history or [])[-10:]:
        role = "Usuario" if h.role == "user" else "Asistente"
        history_lines.append(f"{role}: {h.content}")

    tone = inp.tone or "neutro"
    tone_instruction = "Tono formal y profesional." if tone == "formal" else "Tono neutro y profesional."

    composed = (
        f"{tone_instruction}\n"
        "No inventes datos de personas, empresa, fechas o causas.\n"
        f"HISTORIAL:\n{chr(10).join(history_lines) if history_lines else '(sin historial)'}\n\n"
        f"CONSULTA ACTUAL:\n{question}"
    )

    rag = ask(AskIn(question=composed, collection=inp.collection, top_k=inp.top_k or TOP_K))
    rag_answer = (rag.get("answer") or "").strip()
    sources = rag.get("sources") or []

    best_similarity = _best_similarity(sources)
    sources_count = len(sources)
    answer_len = len(rag_answer or "")
    rag_sufficient_base = (sources_count >= 2 and best_similarity >= 0.72) or (
        answer_len >= 400 and sources_count >= 1
    )

    ql = question.lower()
    needs_official_lookup = any(k in ql for k in [
        "texto exacto",
        "publicado hoy",
        "hoy en el boe",
        "boe de hoy",
        "enlace oficial",
        "enlaces oficiales",
        "dame enlace",
        "dame el enlace",
        "última publicación",
        "ultima publicacion",
        "vigente",
        "vigente hoy",
        "normativa vigente",
        "requisitos",
        "solicitar",
        "trámite",
        "tramite",
        "prestación",
        "prestacion",
        "subsidio",
        "cuantía",
        "cuantia",
        "plazo",
        "documentación",
        "documentacion",
        "donde",
        "enlace",
        "url",
        "sede",
        "boe",
        "sepe",
        "seguridad social",
        "bocm",
        "borm",
    ])

    provided_url_pre = (inp.url or _extract_url(question) or "").strip() or None
    official_id = _detect_official_id(question)
    official_id_url = None
    official_id_kind = None
    if official_id:
        official_id_kind, official_id_value = official_id
        if official_id_kind == "boe":
            official_id_url = f"https://www.boe.es/buscar/doc.php?id={official_id_value}"
        elif official_id_kind == "borme":
            borme_url, _ = _search_borme_id(official_id_value)
            official_id_url = borme_url

    auto_official_triggered = bool(needs_official_lookup and not provided_url_pre)
    internet_effective = bool(inp.internet or auto_official_triggered)

    if internet_effective and auto_official_triggered:
        rag_sufficient = False
        rag_override_reason = "needs_official_lookup_low_similarity"
    else:
        rag_sufficient = rag_sufficient_base
        rag_override_reason = None

    internet_used = False
    internet_reason = "internet_disabled"
    internet_sources: list[dict] = []
    citations: list[dict] = []
    url_used: list[str] | None = None
    answer = rag_answer

    if internet_effective:
        provided_url = provided_url_pre
        if provided_url and not _is_allowed_url(provided_url):
            internet_reason = "blocked_domain"
            answer = (answer + "\n\nBloqueado: solo fuentes oficiales permitidas.").strip()
        elif official_id_url:
            try:
                web_answer, web_sources = _internet_fetch(official_id_url, question)
                internet_used = True
                internet_reason = "direct_official_id"
                internet_sources = web_sources
                citations = [{"n": 1, "url": official_id_url}]
                url_used = [official_id_url]
                answer = (answer + "\n\n" + web_answer + "\n\nFuentes oficiales consultadas:\n[1] " + official_id_url).strip()
            except Exception:
                internet_reason = "no_official_results"
                answer = (answer + "\n\nNo pude recuperar contenido de la referencia oficial indicada.").strip()
        elif provided_url:
            try:
                web_answer, web_sources = _internet_fetch(provided_url, question)
                internet_used = True
                internet_reason = "rag_insufficient"
                internet_sources = web_sources
                citations = [{"n": 1, "url": provided_url}]
                url_used = [provided_url]
                answer = (answer + "\n\n" + web_answer + "\n\nFuentes oficiales consultadas:\n[1] " + provided_url).strip()
            except Exception:
                internet_reason = "no_official_results"
                answer = (answer + "\n\nNo pude recuperar contenido de la URL oficial indicada.").strip()
        elif rag_sufficient:
            internet_reason = "rag_sufficient"
        else:
            area = inp.area or inp.collection
            auto_answer, auto_sources, auto_reason, auto_urls, auto_citations = _search_official_sources(question, area)
            internet_reason = auto_reason
            if auto_sources:
                internet_used = True
                internet_sources = auto_sources
                url_used = auto_urls
                citations = auto_citations
                answer = (answer + "\n\n" + auto_answer).strip()
            else:
                internet_used = False
                answer = (answer + "\n\n" + auto_answer).strip()

    doc_type, template_id = _detect_template(question)
    extracted_fields = _extract_fields(inp.history or [], question)
    if template_id:
        extracted_fields["_template_id"] = template_id

    qdoc = question.lower()
    wants_docx = any(k in qdoc for k in [
        "docx", "word", "archivo", "descargar", "descargable",
        "dámelo en word", "damelo en word", "en word", "en docx", "para descargar", "descarga"
    ])
    wants_blank = any(k in qdoc for k in [
        "en blanco", "plantilla en blanco", "modelo en blanco", "con huecos", "____"
    ])

    required = []
    if template_id in DOCX_TEMPLATES:
        required = DOCX_TEMPLATES[template_id]["required_fields"]
    elif template_id in XLSX_TEMPLATES:
        required = XLSX_TEMPLATES[template_id]["required_fields"]

    missing = [f for f in required if not str(extracted_fields.get(f, "")).strip()]
    needs_data = bool(template_id and missing)
    can_generate_doc = bool(template_id and not missing)
    download_url = None

    if template_id == "carta_despido_objetivo" and wants_docx and wants_blank:
        generated = _generate_docx_file(template_id, {}, None, allow_blank=True)
        download_url = generated["download_url"]
        needs_data = False
        missing = []
        can_generate_doc = True
        doc_type = "docx"
        extracted_fields = {}
        answer = (answer + "\n\nHe generado la plantilla en blanco en Word para descarga.").strip()
    elif template_id == "carta_despido_objetivo" and wants_docx and not wants_blank:
        missing = [
            "empresa_nombre", "empresa_cif", "empresa_domicilio", "trabajador_nombre",
            "trabajador_dni", "fecha_efectos", "causa", "localidad_fecha_firma"
        ]
        needs_data = True
        can_generate_doc = False
        answer = (
            "Resumen de lo entendido:\n"
            "Quieres generar una carta de despido en Word.\n\n"
            "Faltan datos:\n- " + "\n- ".join(missing)
        )
    elif needs_data:
        answer = (
            "Resumen de lo entendido:\n"
            f"Quieres generar una plantilla ({template_id}).\n\n"
            "Datos necesarios:\n- " + "\n- ".join(missing) +
            "\n\nNo voy a inventar datos. Si lo prefieres, te preparo una plantilla en blanco con ____."
        )

    if internet_used and url_used:
        refs = "\n".join([f"[{i}] {u}" for i, u in enumerate(url_used, start=1)])
        if "Fuentes oficiales consultadas:" not in answer:
            answer = (answer + "\n\nFuentes oficiales consultadas:\n" + refs).strip()
    elif not internet_used:
        answer = re.sub(r"https?://\S+", "[enlace no verificado]", answer)

    draft_text = (answer or "").strip() or None
    can_generate_doc_out = bool(draft_text)
    doc_type_out = doc_type if doc_type else ("docx" if can_generate_doc_out else None)

    return {
        "answer": answer,
        "sources": sources,
        "rag_sufficient": rag_sufficient,
        "internet_used": internet_used,
        "internet_reason": internet_reason,
        "internet_sources": internet_sources,
        "citations": citations,
        "internet_effective": internet_effective,
        "auto_official_triggered": auto_official_triggered,
        "rag_override_reason": rag_override_reason,
        "needs_data": needs_data,
        "fields_required": missing,
        "can_generate_doc": can_generate_doc_out,
        "doc_type": doc_type_out,
        "extracted_fields": extracted_fields,
        "draft_text": draft_text,
        "download_url": download_url,
        "url_used": url_used,
    }




def _blank_docx_data(template_id: str) -> dict:
    if template_id == "carta_despido_objetivo":
        return {
            "empresa_nombre": "________",
            "empresa_cif": "________",
            "empresa_domicilio": "________",
            "trabajador_nombre": "________",
            "trabajador_dni": "________",
            "fecha_efectos": "____/____/______",
            "causa": "________________________",
            "localidad_fecha_firma": "____________________, a ____/____/______",
        }
    # fallback genérico
    tpl = DOCX_TEMPLATES.get(template_id) or {}
    req = tpl.get("required_fields", [])
    return {k: "________" for k in req}

def _generate_docx_file(template_id: str, data: dict, draft_text: str | None = None, allow_blank: bool = False) -> dict:
    tpl = DOCX_TEMPLATES.get(template_id)
    if not tpl:
        raise HTTPException(status_code=400, detail="template_id de DOCX no soportado")

    if template_id == "doc_generico":
        body = str((data or {}).get("body", "")).strip()
        if not body:
            raise HTTPException(status_code=422, detail={"needs_data": True, "fields_required": ["body"]})
        safe_data = data or {}
        missing = []
    else:
        missing = [f for f in tpl["required_fields"] if not str((data or {}).get(f, "")).strip()]
        if missing and not allow_blank:
            raise HTTPException(status_code=400, detail={"needs_data": True, "fields_required": missing})
        safe_data = _blank_docx_data(template_id) if allow_blank else (data or {})

    doc = tpl["render"](safe_data, draft_text)
    doc_id = str(uuid.uuid4())
    path = DOCS_DIR / f"{doc_id}.docx"
    doc.save(str(path))
    return {"doc_id": doc_id, "download_url": f"/documents/{doc_id}", "missing": missing}


@app.get("/documents/templates")
def list_document_templates():
    return [
        {
            "id": t["id"],
            "title": t["title"],
            "description": t["description"],
            "required_fields": t["required_fields"],
            "optional_fields": t["optional_fields"],
        }
        for t in DOCX_TEMPLATES.values()
    ]


@app.post("/documents/generate")
def generate_document(inp: DocumentGenerateIn):
    data = inp.data or {}
    blank_mode = (not data) or bool(data.get("blank"))
    generated = _generate_docx_file(inp.template_id, data, inp.draft_text, allow_blank=blank_mode)
    return {"doc_id": generated["doc_id"], "download_url": generated["download_url"]}


@app.get("/documents/{doc_id}")
def download_document(doc_id: str):
    try:
        uuid.UUID(doc_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="doc_id inválido")

    path = DOCS_DIR / f"{doc_id}.docx"
    if not path.exists():
        raise HTTPException(status_code=404, detail="Documento no encontrado")

    return FileResponse(
        str(path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{doc_id}.docx",
    )


@app.get("/spreadsheets/templates")
def list_spreadsheet_templates():
    return [
        {
            "id": t["id"],
            "title": t["title"],
            "description": t["description"],
            "required_fields": t["required_fields"],
            "optional_fields": t["optional_fields"],
        }
        for t in XLSX_TEMPLATES.values()
    ]


@app.post("/spreadsheets/generate")
def generate_spreadsheet(inp: SpreadsheetGenerateIn):
    tpl = XLSX_TEMPLATES.get(inp.template_id)
    if not tpl:
        raise HTTPException(status_code=400, detail="template_id de XLSX no soportado")

    missing = [f for f in tpl["required_fields"] if not str((inp.data or {}).get(f, "")).strip()]
    if missing:
        raise HTTPException(status_code=400, detail={"needs_data": True, "fields_required": missing})

    rows = inp.data.get("rows", []) if isinstance(inp.data, dict) else []
    wb = tpl["render"](inp.data or {}, rows)

    sheet_id = str(uuid.uuid4())
    path = SHEETS_DIR / f"{sheet_id}.xlsx"
    wb.save(str(path))

    return {"sheet_id": sheet_id, "download_url": f"/spreadsheets/{sheet_id}"}


@app.get("/spreadsheets/{sheet_id}")
def download_spreadsheet(sheet_id: str):
    try:
        uuid.UUID(sheet_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="sheet_id inválido")

    path = SHEETS_DIR / f"{sheet_id}.xlsx"
    if not path.exists():
        raise HTTPException(status_code=404, detail="Spreadsheet no encontrado")

    return FileResponse(
        str(path),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        filename=f"{sheet_id}.xlsx",
    )
